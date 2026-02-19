# utils/resume_export.py
"""
Export tailored resume as:
- Formatted .docx (Word document)
- Plain text (for ATS paste)
- JSON (for further processing)
"""

import os
import json
from datetime import datetime
from typing import Dict, Any
from loguru import logger


def export_resume_docx(resume: Dict[str, Any], job_id: str, output_dir: str = "output") -> str:
    """
    Export resume as a formatted Word document (.docx)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        header = resume.get("header", {})

        # ── NAME ──
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(header.get("name", "Your Name"))
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # ── CONTACT ──
        contact_parts = [
            header.get("email", ""),
            header.get("phone", ""),
            header.get("linkedin", ""),
            header.get("github", ""),
        ]
        contact_str = "  |  ".join(p for p in contact_parts if p)
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(contact_str)
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        if header.get("available_from"):
            avail_para = doc.add_paragraph()
            avail_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            avail_run = avail_para.add_run(f"Available from: {header['available_from']}")
            avail_run.italic = True
            avail_run.font.size = Pt(9)
            avail_run.font.color.rgb = RGBColor(0x00, 0x7A, 0x33)

        # ── SUMMARY ──
        _add_section_header(doc, "PROFESSIONAL SUMMARY")
        summary_para = doc.add_paragraph(resume.get("summary", ""))
        summary_para.style.font.size = Pt(10)

        # ── TECHNICAL SKILLS ──
        _add_section_header(doc, "TECHNICAL SKILLS")
        for category, skills in resume.get("technical_skills", {}).items():
            skill_para = doc.add_paragraph()
            bold_run = skill_para.add_run(f"{category}: ")
            bold_run.bold = True
            bold_run.font.size = Pt(10)
            skill_run = skill_para.add_run(", ".join(skills))
            skill_run.font.size = Pt(10)

        # ── PROJECTS ──
        _add_section_header(doc, "PROJECTS")
        for proj in resume.get("projects", []):
            proj_para = doc.add_paragraph()
            name_run = proj_para.add_run(proj.get("name", ""))
            name_run.bold = True
            name_run.font.size = Pt(10)

            tech_run = proj_para.add_run(f"  [{', '.join(proj.get('tech', []))}]")
            tech_run.italic = True
            tech_run.font.size = Pt(9)
            tech_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            desc_para = doc.add_paragraph(proj.get("description", ""), style="List Bullet")
            desc_para.style.font.size = Pt(10)

            if proj.get("metrics"):
                metrics_str = " • ".join(proj["metrics"])
                m_para = doc.add_paragraph(f"Key metrics: {metrics_str}")
                m_para.style.font.size = Pt(9)
                m_para.runs[0].italic = True

        # ── EDUCATION ──
        _add_section_header(doc, "EDUCATION")
        for edu in resume.get("education", []):
            edu_para = doc.add_paragraph()
            deg_run = edu_para.add_run(edu.get("degree", ""))
            deg_run.bold = True
            deg_run.font.size = Pt(10)

            inst_para = doc.add_paragraph(
                f"{edu.get('institution', '')}  |  {edu.get('year', '')}  |  CGPA: {edu.get('cgpa', '')}"
            )
            inst_para.style.font.size = Pt(10)

        # ── CERTIFICATIONS ──
        if resume.get("certifications"):
            _add_section_header(doc, "CERTIFICATIONS")
            for cert in resume["certifications"]:
                cert_para = doc.add_paragraph(cert, style="List Bullet")
                cert_para.style.font.size = Pt(10)

        # ── ACHIEVEMENTS ──
        if resume.get("achievements"):
            _add_section_header(doc, "ACHIEVEMENTS")
            for ach in resume["achievements"]:
                ach_para = doc.add_paragraph(ach, style="List Bullet")
                ach_para.style.font.size = Pt(10)

        # Save
        os.makedirs(output_dir, exist_ok=True)
        filename = f"resume_{job_id}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        logger.info(f"✅ Resume exported: {filepath}")
        return filepath

    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"Resume export failed: {e}")
        return ""


def _add_section_header(doc, title: str):
    """Add a styled section header with bottom border"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8E)

    # Add bottom border via XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A8E")
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()  # Spacing


def export_resume_text(resume: Dict[str, Any]) -> str:
    """Export resume as plain text (ATS-safe)"""
    lines = []
    header = resume.get("header", {})

    lines.append(header.get("name", "").upper())
    lines.append(f"{header.get('email')} | {header.get('phone')} | {header.get('linkedin')}")
    lines.append(f"GitHub: {header.get('github')}")
    lines.append(f"Available from: {header.get('available_from')}")
    lines.append("")

    lines.append("SUMMARY")
    lines.append("-" * 40)
    lines.append(resume.get("summary", ""))
    lines.append("")

    lines.append("TECHNICAL SKILLS")
    lines.append("-" * 40)
    for cat, skills in resume.get("technical_skills", {}).items():
        lines.append(f"{cat}: {', '.join(skills)}")
    lines.append("")

    lines.append("PROJECTS")
    lines.append("-" * 40)
    for proj in resume.get("projects", []):
        lines.append(f"\n{proj.get('name', '')}")
        lines.append(f"Tech: {', '.join(proj.get('tech', []))}")
        lines.append(proj.get("description", ""))
        if proj.get("metrics"):
            lines.append("Metrics: " + " | ".join(proj["metrics"]))
    lines.append("")

    lines.append("EDUCATION")
    lines.append("-" * 40)
    for edu in resume.get("education", []):
        lines.append(f"{edu.get('degree')} | {edu.get('institution')} | {edu.get('year')} | CGPA: {edu.get('cgpa')}")
    lines.append("")

    lines.append("CERTIFICATIONS")
    lines.append("-" * 40)
    for cert in resume.get("certifications", []):
        lines.append(f"• {cert}")

    return "\n".join(lines)
